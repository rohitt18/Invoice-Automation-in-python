import os
import subprocess
import datetime as dt
import tkinter as tk
from tkinter import filedialog, StringVar, OptionMenu
from tkinter import messagebox

from docx import Document
import docx

from docx2pdf import convert

replacements = {}

class InvoiceAutomation:
    replacements = {}
    def __init__(self):
        self.root = tk.Tk()
        self.root.title('Invoice Automation')
        self.root.geometry('500x600')

        self.partner_label = tk.Label(self.root, text='Partner')
        self.partner_street_label = tk.Label(self.root, text='Partner Street')
        self.partner_zip_city_country_label = tk.Label(self.root, text='Partner ZIP CITY COUNTRY')
        self.invoice_number_label = tk.Label(self.root, text='Invoice Numebr')
        self.service_description_label = tk.Label(self.root, text='Service Description')
        self.service_amount_label = tk.Label(self.root, text='Service Amount')
        self.service_single_price_label = tk.Label(self.root, text='Service Single Price')
        self.payment_method_label = tk.Label(self.root, text='Payment Method')

        self.payment_methods = {
            'Main Bank':{
                'Recipient': 'Infinite Software Company',
                'Bank': 'Bank Of India',
                'Account Number': '324325891589',
                'IFSC': 'BKID12345678'
            },
            'Second Bank': {
                'Recipient': 'Infinite Software Company',
                'Bank': 'State Bank Of India',
                'Account Number': '668812574422',
                'IFSC': 'SBI000126891'
            },
            'Private Bank': {
                'Recipient': 'Rohit Gupta',
                'Bank': 'Kotak Mahindra Bank',
                'Account Number': '001166882191',
                'IFSC': 'KBI012680222'
            }
        }

        self.partner_entry= tk.Entry(self.root)
        self.partner_street_entry = tk.Entry(self.root)
        self.partner_zip_city_country_entry = tk.Entry(self.root)
        self.invoice_number_entry = tk.Entry(self.root)
        self.service_description_entry = tk.Entry(self.root)
        self.service_amount_entry = tk.Entry(self.root)
        self.service_single_price_entry = tk.Entry(self.root)

        payment_methods = ["Main Bank", "Second Bank", "Private Bank"]

        self.payment_method = StringVar()
        self.payment_method.set(payment_methods[0])  # Set the default value


        self.payment_method.dropdown = tk.OptionMenu(self.root, self.payment_method,*payment_methods )


        self.create_button = tk.Button(self.root, text="Create Invoice", command = self.create_invoice)

        padding_options = {'fill': 'x', 'expand':True, 'padx':5, 'pady': 2}

        self.partner_label.pack(padding_options)
        self.partner_entry.pack(padding_options)
        self.partner_street_label.pack(padding_options)
        self.partner_street_entry.pack(padding_options)
        self.partner_zip_city_country_label.pack(padding_options)
        self.partner_zip_city_country_entry.pack(padding_options)
        self.invoice_number_label.pack(padding_options)
        self.invoice_number_entry.pack(padding_options)
        self.service_description_label.pack(padding_options)
        self.service_description_entry.pack(padding_options)
        self.service_amount_label.pack(padding_options)
        self.service_amount_entry.pack(padding_options)
        self.service_single_price_label.pack(padding_options)
        self.service_single_price_entry.pack(padding_options)
        self.payment_method_label.pack(padding_options)
        self.payment_method.dropdown.pack(padding_options)
        self.create_button.pack(padding_options)

        print("Starting mainloop")
        self.root.mainloop()

    @staticmethod
    def replace_text(paragraph, old_text, new_text):
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)


    def create_invoice(self):
        doc = docx.Document('template.docx')

        selected_payment_method = self.payment_methods[self.payment_method.get()]

        try:
            global replacements
            replacements = {
                    "[Date]": dt.datetime.today().strftime('%d-%m-%Y'),
                    "[Partner]": self.partner_entry.get(),
                    "[Partner Street]": self.partner_street_entry.get(),
                    "[Partner_ZIP_City_Country]": self.partner_zip_city_country_entry.get(),
                    "[Invoice Number]": self.invoice_number_entry.get(),
                    "[Service Description]": self.service_description_entry.get(),
                    "[Amount]": self.service_amount_entry.get(),
                    "[Single Price]": f"Rs{float(self.service_single_price_entry.get()):.2f}",
                    "[Full Price]": f'Rs{float(self.service_amount_entry.get()) * float(self.service_single_price_entry.get()):.2f}',
                    "[Recipient]": selected_payment_method['Recipient'],
                    "[Bank]": selected_payment_method['Bank'],
                    "[Account Number]": selected_payment_method['Account Number'],
                    "[IFSC]": selected_payment_method['IFSC'],
                }

        except ValueError:
            messagebox.showerror(title='Error',message='Invalid amount or price')
            return

        for paragraph in list(doc.paragraphs):
            for old_text, new_text in replacements.items():
                self.replace_text(paragraph, old_text, new_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            self.replace_text(paragraph, old_text, new_text)

        save_path = filedialog.asksaveasfilename(defaultextension='.pdf', filetypes=[('PDF documents', '*.pdf')])

        doc.save('filled.docx')

        convert('filled.docx', save_path)

        messagebox.showinfo(title='success', message='Invoice created and saved successfully ')


if __name__ == '__main__':
    InvoiceAutomation()

